from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "crypto_pulse_db_report.docx"
BLACK = (0, 0, 0)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for idx, width in enumerate(widths):
        tbl_grid.gridCol_lst[idx].set(qn("w:w"), str(width))

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*BLACK)
    run.bold = bold


def add_paragraph(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        style_run(run, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        style_run(rest)
    else:
        run = p.add_run(text)
        style_run(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    style_run(run)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    style_run(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    style_run(run, size=16, bold=True)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=14)
    return p


def add_table(doc):
    add_heading(doc, "3. Які теми лабораторних робіт використані у проєкті", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 3300, 3860])
    headers = ["Лабораторна робота", "Що використовується у проєкті", "Де це видно"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        style_run(run, bold=True)
    rows = [
        (
            "ЛР 1. Створення схеми БД",
            "Таблиці, первинні ключі, зовнішні ключі, унікальні обмеження, NOT NULL, DEFAULT, сурогатні ключі, індекси",
            "У SQL-схемі таблиць users, user_favorites, alerts, ml_models, model_predictions",
        ),
        (
            "ЛР 2. Написання SQL-запитів",
            "SELECT, INSERT, UPDATE, DELETE, WHERE, GROUP BY, ORDER BY, агрегатні функції, обчислювані поля",
            "У файлах bot.py, ml_engine.py та в SQL-скрипті структури БД",
        ),
        (
            "ЛР 3. Підпрограми СУБД",
            "Функції СУБД для автоматизації логіки користувачів",
            "У SQL-функціях handle_new_user(), check_email_exists(), delete_current_user()",
        ),
        (
            "ЛР 4. Тригери",
            "Автоматичне виконання дії після події в БД",
            "У тригері on_auth_user_created",
        ),
        (
            "ЛР 5. Адміністрування БД",
            "Права доступу, ролі, GRANT, RLS-політики, контроль доступу до рядків",
            "У SQL-блоках GRANT, ENABLE ROW LEVEL SECURITY, CREATE POLICY",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = ""
            run = cells[idx].paragraphs[0].add_run(text)
            style_run(run)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(*BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.0

    title_top_1 = doc.add_paragraph()
    title_top_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title_top_1.add_run("Міністерство освіти і науки України"), size=14, bold=True)

    title_top_2 = doc.add_paragraph()
    title_top_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title_top_2.add_run("[Назва закладу вищої освіти]"), size=14, bold=True)

    title_top_3 = doc.add_paragraph()
    title_top_3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title_top_3.add_run("[Назва кафедри]"), size=14, bold=True)

    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title.add_run("ЗВІТ"), size=16, bold=True)

    title_subject = doc.add_paragraph()
    title_subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title_subject.add_run("з дисципліни «Бази даних»"), size=14)

    title_project = doc.add_paragraph()
    title_project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title_project.add_run("за проєктом Crypto Pulse"), size=14, bold=True)

    for _ in range(4):
        doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(author.add_run("Виконав: студент групи [група]\n[ПІБ студента]"), size=14)

    reviewer = doc.add_paragraph()
    reviewer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(reviewer.add_run("Перевірив: [ПІБ викладача]"), size=14)

    for _ in range(5):
        doc.add_paragraph()

    city_year = doc.add_paragraph()
    city_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(city_year.add_run("[Місто] - 2026"), size=14)

    doc.add_page_break()

    title_main = doc.add_paragraph()
    title_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title_main.add_run("Звіт з баз даних за проєктом Crypto Pulse"), size=16, bold=True)

    add_heading(doc, "Мета роботи", level=1)
    add_paragraph(
        doc,
        "Метою роботи є аналіз бази даних проєкту Crypto Pulse, визначення її структури, основних таблиць, "
        "зв'язків між ними, а також встановлення того, які теми лабораторних робіт з дисципліни «Бази даних» уже реалізовані в проєкті. "
        "Окремою метою є пояснення використаних SQL-конструкцій, функцій СУБД, тригерів, прав доступу та запланованих механізмів "
        "розмежування функціональності за рівнями підписки.",
    )

    add_heading(doc, "Хід роботи", level=1)
    for item in [
        "Проаналізовано SQL-схему бази даних проєкту.",
        "Визначено основні таблиці та зв'язки між ними.",
        "Розглянуто SQL-запити, що використовуються в Python-коді проєкту.",
        "Проаналізовано функції СУБД, які відповідають за роботу з користувачами.",
        "Розглянуто тригер автоматичного створення профілю користувача.",
        "Досліджено механізми прав доступу та Row-Level Security.",
        "Сформульовано запланований розвиток системи підписок і обмеження можливостей користувачів за тарифами.",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "1. Загальна характеристика проєкту та бази даних", level=1)
    add_paragraph(
        doc,
        "Проєкт Crypto Pulse є інформаційною системою для роботи з криптовалютними даними. "
        "У системі користувач може зареєструватися, мати власний профіль, додавати монети до списку обраних, "
        "створювати цінові сповіщення, а також переглядати результати роботи ML-моделі, яка формує торгові сигнали.",
    )
    add_paragraph(
        doc,
        "Для зберігання даних у проєкті використовується реляційна база даних PostgreSQL у середовищі Supabase. "
        "Supabase було обрано тому, що це хмарна платформа, яка надає готову PostgreSQL-базу, засоби автентифікації користувачів, "
        "керування доступом і зручний веб-інтерфейс для роботи з таблицями. Завдяки розміщенню БД у спільному хмарному середовищі "
        "всі учасники команди та всі частини застосунку працюють з одними й тими самими актуальними даними, а не з різними локальними копіями. "
        "Це спрощує спільну розробку, тестування та підтримку проєкту. Крім того, Supabase поєднує PostgreSQL із системою авторизації "
        "та механізмом Row-Level Security, тому в проєкті можна не лише зберігати дані, а й обмежувати доступ користувачів до власних записів. "
        "У Supabase окремо існує системна таблиця auth.users, яка відповідає за автентифікацію користувачів. "
        "Для прикладних даних проєкту використовується власна схема public.",
    )
    add_paragraph(doc, "Основні таблиці бази даних:")
    for item in [
        "auth.users - системні облікові записи користувачів Supabase.",
        "public.users - профілі користувачів застосунку.",
        "public.user_favorites - криптовалюти, які користувач додав до обраного.",
        "public.alerts - персональні сповіщення про досягнення певної ціни.",
        "public.ml_models - натреновані ML-моделі для різних криптовалют і таймфреймів.",
        "public.model_predictions - історія прогнозів, які сформувала ML-модель.",
    ]:
        add_numbered(doc, item)
    add_paragraph(doc, "Основні зв'язки між таблицями:")
    for item in [
        "public.users.id пов'язаний із auth.users.id.",
        "public.user_favorites.user_id посилається на auth.users.id.",
        "public.alerts.user_id посилається на auth.users.id.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Схема бази даних проєкту", level=1)
    add_paragraph(
        doc,
        "Нижче наведено візуальну ER-схему основних таблиць бази даних проєкту.",
    )
    doc.add_picture(str(ROOT / "crypto_pulse_schema.png"), width=Inches(6.5))
    add_paragraph(
        doc,
        "На схемі видно основні таблиці, їх поля та зв'язки між ними. Центральною таблицею профілів є users. "
        "Із нею пов'язані таблиці alerts і user_favorites, оскільки вони містять дані, що належать конкретному користувачу. "
        "Таблиці ml_models і model_predictions використовуються для роботи машинного навчання: перша зберігає моделі, а друга - результати їх прогнозів.",
    )

    add_table(doc)

    add_heading(doc, "4. Лабораторна робота №1. Створення схеми БД", level=1)
    add_heading(doc, "4.1. Що саме з ЛР 1 використано у проєкті", level=2)
    add_paragraph(
        doc,
        "У межах першої лабораторної роботи потрібно побудувати схему бази даних, визначити таблиці, поля, типи даних, "
        "ключі та обмеження цілісності. У проєкті Crypto Pulse це реалізовано через SQL-схему бази даних.",
    )
    for item in [
        "кілька взаємопов'язаних таблиць;",
        "первинні ключі;",
        "зовнішні ключі;",
        "сурогатні ключі;",
        "унікальні обмеження;",
        "обмеження NOT NULL;",
        "значення за замовчуванням DEFAULT;",
        "індекси для швидкого пошуку;",
        "каскадне видалення пов'язаних записів.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.2. Які таблиці створені і для чого вони потрібні", level=2)
    lab1_blocks = [
        (
            "Таблиця public.users",
            "Таблиця users зберігає додаткову інформацію про користувача, якої немає в системній таблиці Supabase Auth: "
            "ім'я, прізвище, username, телефон, регіон, аватар, тарифний план, Telegram-дані та службові дати.",
            "CREATE TABLE IF NOT EXISTS public.users (\n"
            "  id UUID PRIMARY KEY REFERENCES auth.users(id) ON DELETE CASCADE,\n"
            "  email TEXT NOT NULL UNIQUE,\n"
            "  is_active BOOLEAN NOT NULL DEFAULT true,\n"
            "  username VARCHAR(100),\n"
            "  first_name VARCHAR(100),\n"
            "  last_name VARCHAR(100)\n"
            ");",
            [
                "id є первинним ключем і одночасно зовнішнім ключем на auth.users(id).",
                "ON DELETE CASCADE автоматично видаляє профіль після видалення облікового запису.",
                "email має обмеження UNIQUE, тому дублікати неможливі.",
                "is_active має значення за замовчуванням true.",
            ],
        ),
        (
            "Таблиця public.user_favorites",
            "Таблиця user_favorites зберігає монети, які користувач додав до обраного.",
            "CREATE TABLE IF NOT EXISTS public.user_favorites (\n"
            "  id BIGSERIAL PRIMARY KEY,\n"
            "  user_id UUID NOT NULL REFERENCES auth.users(id) ON DELETE CASCADE,\n"
            "  coin_id TEXT NOT NULL,\n"
            "  symbol VARCHAR(20) NOT NULL,\n"
            "  name TEXT NOT NULL,\n"
            "  created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),\n"
            "  CONSTRAINT user_favorites_user_symbol_unique UNIQUE (user_id, symbol)\n"
            ");",
            [
                "BIGSERIAL створює сурогатний ключ автоматично.",
                "user_id формує зв'язок один користувач - багато обраних монет.",
                "UNIQUE (user_id, symbol) не дозволяє додати одну монету двічі.",
            ],
        ),
        (
            "Таблиця public.alerts",
            "Таблиця alerts містить персональні умови сповіщення за ціною криптовалюти.",
            "CREATE TABLE IF NOT EXISTS public.alerts (\n"
            "  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
            "  user_id UUID REFERENCES auth.users(id),\n"
            "  symbol TEXT NOT NULL,\n"
            "  condition TEXT NOT NULL,\n"
            "  target_price NUMERIC NOT NULL,\n"
            "  is_active BOOLEAN DEFAULT true\n"
            ");",
            [
                "NUMERIC використано для точного зберігання ціни.",
                "Поле is_active описує, чи сповіщення ще активне.",
                "user_id закріплює кожне сповіщення за конкретним користувачем.",
            ],
        ),
        (
            "Таблиця public.ml_models",
            "Таблиця ml_models зберігає натреновані моделі машинного навчання.",
            "CREATE TABLE IF NOT EXISTS public.ml_models (\n"
            "  id INTEGER PRIMARY KEY DEFAULT nextval('ml_models_id_seq'::regclass),\n"
            "  symbol VARCHAR NOT NULL,\n"
            "  interval VARCHAR NOT NULL,\n"
            "  model_binary BYTEA NOT NULL,\n"
            "  accuracy REAL,\n"
            "  features ARRAY,\n"
            "  trained_at TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP\n"
            ");",
            [
                "BYTEA дозволяє зберігати модель у бінарному вигляді.",
                "ARRAY зберігає список ознак, на яких навчалась модель.",
                "trained_at дозволяє визначити, коли модель була навчена.",
            ],
        ),
        (
            "Таблиця public.model_predictions",
            "Таблиця model_predictions зберігає історію прогнозів моделі.",
            "CREATE TABLE IF NOT EXISTS public.model_predictions (\n"
            "  id INTEGER PRIMARY KEY DEFAULT nextval('model_predictions_id_seq'::regclass),\n"
            "  symbol VARCHAR NOT NULL,\n"
            "  interval VARCHAR NOT NULL,\n"
            "  signal VARCHAR NOT NULL,\n"
            "  price NUMERIC NOT NULL,\n"
            "  confidence NUMERIC,\n"
            "  accuracy NUMERIC,\n"
            "  raw_prediction VARCHAR,\n"
            "  stop_loss NUMERIC,\n"
            "  take_profit NUMERIC,\n"
            "  created_at TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP\n"
            ");",
            [
                "У цю таблицю записуються реальні результати роботи ML-моделі.",
                "Саме її читає бот, коли перевіряє свіжість прогнозів.",
            ],
        ),
    ]
    for title_text, paragraph, code, bullets in lab1_blocks:
        add_heading(doc, title_text, level=3)
        add_paragraph(doc, paragraph)
        add_code_block(doc, code)
        if title_text == "Таблиця public.users":
            add_paragraph(
                doc,
                "Оператор CREATE TABLE створює таблицю public.users, якщо вона ще не існує. "
                "У ній додаються поля для ідентифікатора користувача, email, стану акаунта та персональних даних. "
                "Одразу під час створення таблиці задаються первинний ключ, зовнішній ключ, унікальність email та значення за замовчуванням.",
            )
        elif title_text == "Таблиця public.user_favorites":
            add_paragraph(
                doc,
                "Через CREATE TABLE створюється окрема таблиця для списку обраних монет. "
                "До неї додаються id запису, id користувача, технічний id монети, її символ, назва та дата створення. "
                "Обмеження UNIQUE (user_id, symbol) задає правило, за яким одна й та сама монета не може двічі зберігатися в обраному одного користувача.",
            )
        elif title_text == "Таблиця public.alerts":
            add_paragraph(
                doc,
                "Цей блок створює таблицю цінових сповіщень. "
                "У стовпцях зберігаються користувач, монета, цільова ціна, умова спрацювання, активність і факт спрацювання сповіщення. "
                "Таким чином таблиця описує і саме сповіщення, і його поточний стан.",
            )
        elif title_text == "Таблиця public.ml_models":
            add_paragraph(
                doc,
                "Цей CREATE TABLE створює таблицю для збереження навчених ML-моделей. "
                "У ній передбачені поля для монети, таймфрейму, бінарного вмісту моделі, точності, списку ознак і часу навчання. "
                "Поле trained_at потрібне для перевірки актуальності моделі.",
            )
        elif title_text == "Таблиця public.model_predictions":
            add_paragraph(
                doc,
                "Тут створюється таблиця журналу прогнозів. "
                "До неї додаються поля для монети, таймфрейму, сигналу, ціни, впевненості моделі та часу створення запису. "
                "Це дає змогу не лише показати прогноз користувачу, а й зберегти його в історії для подальшого аналізу.",
            )
        for item in bullets:
            add_bullet(doc, item)
    add_heading(doc, "4.3. Де це видно в коді проєкту", level=2)
    add_paragraph(
        doc,
        "У файлі ml_engine.py функція save_model_to_db() записує нові моделі в таблицю ml_models, "
        "а функція log_prediction_to_db() додає нові записи до model_predictions. У файлі bot.py функція "
        "get_prediction_age_map() читає історію прогнозів із model_predictions. Це показує, що таблиці не лише створені, "
        "а реально використовуються прикладним кодом.",
    )
    add_heading(doc, "4.4. Висновок по ЛР 1", level=2)
    add_paragraph(
        doc,
        "У проєкті використано повну реляційну схему даних із різними типами ключів, обмеженнями цілісності та індексами. "
        "Це є прямим практичним застосуванням тем першої лабораторної роботи.",
    )

    add_heading(doc, "5. Лабораторна робота №2. Написання SQL-запитів", level=1)
    add_heading(doc, "5.1. Що саме з ЛР 2 використано у проєкті", level=2)
    add_paragraph(
        doc,
        "У другій лабораторній роботі вивчаються різні типи SQL-запитів. У проєкті Crypto Pulse вони виконуються "
        "безпосередньо з Python-коду через бібліотеку psycopg2.",
    )
    for item in [
        "SELECT;",
        "INSERT;",
        "UPDATE;",
        "DELETE;",
        "WHERE;",
        "AND;",
        "MAX;",
        "GROUP BY;",
        "ORDER BY;",
        "IS NOT NULL;",
        "ILIKE;",
        "обчислюване поле в результаті SELECT.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.2. Приклад SELECT із групуванням у bot.py", level=2)
    add_code_block(
        doc,
        "SELECT\n"
        "    symbol,\n"
        "    interval,\n"
        "    EXTRACT(EPOCH FROM (NOW() - MAX(created_at))) / 3600.0 AS age_hours\n"
        "FROM model_predictions\n"
        "GROUP BY symbol, interval;",
    )
    add_paragraph(
        doc,
        "Оператор SELECT вибирає символ монети, таймфрейм і додатково рахує поле age_hours. "
        "Функція MAX(created_at) знаходить найновіший прогноз у кожній групі, а GROUP BY symbol, interval формує окремий результат "
        "для кожної монети й кожного таймфрейму.",
    )
    for item in [
        "знаходить усі прогнози в model_predictions;",
        "через MAX(created_at) обирає останній прогноз для кожної пари symbol + interval;",
        "обчислює, скільки годин минуло від останнього прогнозу;",
        "через GROUP BY формує окремий результат для кожної монети й таймфрейму.",
    ]:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "Практичний сенс запиту полягає в тому, що бот не перераховує прогноз без потреби, якщо наявний прогноз ще свіжий.",
    )
    add_heading(doc, "5.3. Приклад INSERT у ml_engine.py", level=2)
    add_code_block(
        doc,
        "INSERT INTO model_predictions\n"
        "(symbol, interval, signal, price, confidence, accuracy, raw_prediction, stop_loss, take_profit, created_at)\n"
        "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s);",
    )
    add_paragraph(
        doc,
        "Оператор INSERT INTO додає новий рядок до таблиці model_predictions. "
        "Після назви таблиці перелічуються стовпці, у які треба записати дані, а VALUES отримує конкретні значення з Python-коду. "
        "Так кожен прогноз моделі фіксується в історії бази даних.",
    )
    add_paragraph(
        doc,
        "Цей запит використовується у функції log_prediction_to_db() і зберігає кожен сформований прогноз у базі даних. "
        "До таблиці потрапляють монета, таймфрейм, сигнал, ціна, впевненість моделі, точність, рівні stop loss і take profit та час створення.",
    )
    add_heading(doc, "5.4. Приклад DELETE та INSERT для оновлення моделей у ml_engine.py", level=2)
    add_code_block(doc, "DELETE FROM ml_models\nWHERE symbol = %s AND interval = %s;")
    add_paragraph(
        doc,
        "Оператор DELETE видаляє не всі моделі, а лише старий запис для конкретної пари symbol та interval. "
        "Тобто видаляється тільки попередня модель для потрібної монети й таймфрейму.",
    )
    add_code_block(
        doc,
        "INSERT INTO ml_models\n"
        "(symbol, interval, model_binary, accuracy, features, trained_at)\n"
        "VALUES (%s, %s, %s, %s, %s, %s);",
    )
    add_paragraph(
        doc,
        "Після видалення старої моделі INSERT INTO додає нову модель. "
        "До таблиці записуються символ монети, таймфрейм, бінарне представлення моделі, точність, список ознак і час навчання. "
        "Разом ці два запити реалізують заміну старої моделі новою.",
    )
    add_paragraph(
        doc,
        "Функція save_model_to_db() спочатку видаляє стару модель для конкретної монети й таймфрейму, а потім записує нову. "
        "Таким чином у таблиці зберігається актуальна версія моделі без дублювання. "
        "У поточній реалізації це правило підтримується самим прикладним кодом.",
    )
    add_heading(doc, "5.5. Приклад складного SELECT у ml_engine.py", level=2)
    add_code_block(
        doc,
        "SELECT signal, price\n"
        "FROM model_predictions\n"
        "WHERE UPPER(symbol) = %s\n"
        "  AND interval = %s\n"
        "  AND price IS NOT NULL\n"
        "  AND signal IS NOT NULL\n"
        "  AND signal NOT ILIKE 'NO TRADE%'\n"
        "  AND created_at <= NOW() - (%s * INTERVAL '1 hour')\n"
        "ORDER BY created_at DESC\n"
        "LIMIT %s;",
    )
    add_paragraph(
        doc,
        "Цей SELECT вибирає лише ті прогнози, які придатні для перевірки результативності. "
        "Умови WHERE відбирають потрібну монету й таймфрейм, відсікають порожні значення та виключають сигнали NO TRADE. "
        "ORDER BY сортує записи від новіших до старіших, а LIMIT обмежує кількість рядків.",
    )
    add_paragraph(
        doc,
        "Цей запит у функції get_recent_signal_performance() вибирає лише корисні для аналізу прогнози: для потрібної монети, "
        "потрібного таймфрейму, без порожніх значень і без сигналів NO TRADE. Після сортування від найновіших до старіших Python-код "
        "перевіряє, чи справдилися попередні сигнали, і рахує статистику результативності.",
    )
    add_heading(doc, "5.6. Приклад UPDATE у SQL-схемі", level=2)
    add_code_block(
        doc,
        "UPDATE public.users\n"
        "SET active_plan = COALESCE(active_plan, subscription, 'free'),\n"
        "    subscription = COALESCE(subscription, active_plan, 'free'),\n"
        "    billing_cycle = COALESCE(billing_cycle, 'monthly');",
    )
    add_paragraph(
        doc,
        "Оператор UPDATE змінює вже наявні записи в public.users. "
        "Функція COALESCE бере перше ненульове значення, тому старі профілі отримують коректний тариф і тип оплати навіть тоді, "
        "коли частина полів раніше була порожньою.",
    )
    add_paragraph(
        doc,
        "Цей запит приводить старі профілі користувачів до єдиного формату після зміни структури таблиці users.",
    )
    add_heading(doc, "5.7. Фрагменти Python-коду, у яких база даних використовується в роботі системи", level=2)
    add_paragraph(
        doc,
        "Окрім самих SQL-запитів, у проєкті є програмний код, який показує, як застосунок підключається до БД, "
        "записує результати, читає моделі та використовує дані з таблиць для прийняття рішень.",
    )
    python_blocks = [
        (
            "Підключення до бази даних у ml_engine.py",
            "def get_db_connection():\n"
            "    if HAS_STREAMLIT:\n"
            "        try:\n"
            "            return psycopg2.connect(...)\n"
            "        except Exception:\n"
            "            pass\n\n"
            "    return psycopg2.connect(\n"
            "        dbname=os.getenv('POSTGRES_DB'),\n"
            "        user=os.getenv('POSTGRES_USER'),\n"
            "        password=os.getenv('POSTGRES_PASSWORD'),\n"
            "        host=os.getenv('DB_HOST'),\n"
            "        port=os.getenv('DB_PORT', '5432')\n"
            "    )",
            "Функція get_db_connection() створює з'єднання з PostgreSQL через бібліотеку psycopg2. "
            "Якщо застосунок запущений у Streamlit, дані для підключення беруться зі st.secrets, а якщо код працює як бот на сервері - зі змінних оточення. "
            "Це дозволяє одному коду працювати в різних середовищах без зміни логіки БД. "
            "Додатково такий підхід не вимагає записувати логін, пароль і адресу сервера безпосередньо в коді, тому конфіденційні параметри зберігаються окремо від програми.",
        ),
        (
            "Перевірка свіжості прогнозів у bot.py",
            "def get_prediction_age_map():\n"
            "    conn = get_db_connection()\n"
            "    cur = conn.cursor()\n"
            "    cur.execute('''\n"
            "        SELECT symbol, interval,\n"
            "               EXTRACT(EPOCH FROM (NOW() - MAX(created_at))) / 3600.0 AS age_hours\n"
            "        FROM model_predictions\n"
            "        GROUP BY symbol, interval\n"
            "    ''')\n"
            "    rows = cur.fetchall()\n"
            "    return {(symbol, interval): age_hours for symbol, interval, age_hours in rows}",
            "Функція виконує SQL-запит до model_predictions, отримує результати через fetchall() "
            "і перетворює їх на словник Python. Ключем стає пара symbol + interval, а значенням - вік останнього прогнозу в годинах. "
            "Далі ці дані використовуються ботом, щоб не перераховувати занадто свіжі прогнози. "
            "Блок finally гарантує закриття підключення навіть у разі помилки, щоб не витрачати зайві ресурси БД.",
        ),
        (
            "Побудова списку завдань на основі даних із БД у bot.py",
            "def build_smart_jobs(coins):\n"
            "    age_map = get_prediction_age_map()\n"
            "    jobs = []\n"
            "    for coin in coins:\n"
            "        for interval in INTERVALS:\n"
            "            age = age_map.get((coin.upper(), interval))\n"
            "            if age is not None and age < PREDICTION_FRESH_HOURS:\n"
            "                continue\n"
            "            jobs.append({'coin': coin, 'interval': interval})",
            "Тут використовується результат попереднього SQL-запиту. "
            "Функція формує лише ті завдання, які справді потрібно виконати, і пропускає прогнози, що ще залишаються свіжими. "
            "Отже, дані з БД прямо впливають на планування роботи бота. "
            "Прапорець missing показує, чи прогноз відсутній повністю, а велике значення age допомагає підняти такі завдання в пріоритеті.",
        ),
        (
            "Запис нового прогнозу в БД у ml_engine.py",
            "def log_prediction_to_db(data: dict):\n"
            "    conn = get_db_connection()\n"
            "    cursor = conn.cursor()\n"
            "    query = '''\n"
            "        INSERT INTO model_predictions\n"
            "        (symbol, interval, signal, price, confidence, accuracy,\n"
            "         raw_prediction, stop_loss, take_profit, created_at)\n"
            "        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)\n"
            "    '''\n"
            "    cursor.execute(query, values)\n"
            "    conn.commit()",
            "Функція log_prediction_to_db() формує SQL-запит INSERT і через cursor.execute() "
            "додає новий рядок до model_predictions. Після conn.commit() зміни фіксуються в базі, тому кожен сигнал залишається в історії. "
            "Плейсхолдери %s означають, що використано параметризований запит, а допоміжна функція safe_float() не дає записати некоректні числові значення.",
        ),
        (
            "Збереження моделі у БД у ml_engine.py",
            "def save_model_to_db(symbol, interval, model, accuracy, features):\n"
            "    model_bytes = pickle.dumps(model)\n"
            "    cur.execute('DELETE FROM ml_models WHERE symbol = %s AND interval = %s', ...)\n"
            "    cur.execute('''\n"
            "        INSERT INTO ml_models\n"
            "        (symbol, interval, model_binary, accuracy, features, trained_at)\n"
            "        VALUES (%s, %s, %s, %s, %s, %s)\n"
            "    ''', ...)\n"
            "    conn.commit()",
            "Перед записом нової моделі вона перетворюється на бінарний формат через pickle.dumps(). "
            "Потім стара модель для тієї самої пари symbol + interval видаляється, а нова записується в ml_models. "
            "Це дозволяє повторно використовувати вже навчені моделі. "
            "Список features зберігається разом із моделлю, щоб під час наступного використання побудувати вхідні дані в тому самому форматі.",
        ),
        (
            "Завантаження вже навченої моделі з БД у ml_engine.py",
            "def load_model_from_db(symbol, interval):\n"
            "    cur.execute(\n"
            "        'SELECT model_binary, accuracy, features, trained_at '\n"
            "        'FROM ml_models WHERE symbol = %s AND interval = %s',\n"
            "        (symbol, interval)\n"
            "    )\n"
            "    row = cur.fetchone()\n"
            "    if row:\n"
            "        return pickle.loads(row[0]), row[1], row[2]",
            "Функція шукає вже збережену модель у таблиці ml_models і відновлює її через pickle.loads(). "
            "Завдяки цьому модель не потрібно навчати заново при кожному зверненні до системи. "
            "Перевірка trained_at потрібна, щоб не використовувати застарілу модель.",
        ),
        (
            "Аналіз минулих прогнозів у ml_engine.py",
            "def get_recent_signal_performance(...):\n"
            "    cur.execute('''\n"
            "        SELECT signal, price\n"
            "        FROM model_predictions\n"
            "        WHERE UPPER(symbol) = %s\n"
            "          AND interval = %s\n"
            "          AND signal NOT ILIKE 'NO TRADE%%'\n"
            "        ORDER BY created_at DESC\n"
            "        LIMIT %s\n"
            "    ''', ...)",
            "Функція читає попередні прогнози з model_predictions, після чого Python-код перевіряє, "
            "чи спрацювали сигнали LONG або SHORT. Отримана статистика використовується для адаптивного коригування порогів моделі. "
            "Часовий горизонт дозволяє оцінювати прогноз лише після того, як минув достатній проміжок часу для відповідного таймфрейму.",
        ),
        (
            "Запуск циклу бота у bot.py",
            "def run_bot():\n"
            "    coins_to_analyze = get_top_125_coins()\n"
            "    jobs = build_smart_jobs(coins_to_analyze)\n"
            "    for job in jobs:\n"
            "        coin = job['coin']\n"
            "        interval = job['interval']\n"
            "        result = get_ml_signal(coin, interval)",
            "Бот спочатку отримує список монет, потім будує завдання з урахуванням уже збережених у БД прогнозів, "
            "а далі запускає побудову нових сигналів. Після формування прогнозу результат знову записується в БД, тобто база бере участь у всьому циклі роботи системи. "
            "Отже, БД є не лише сховищем, а й джерелом інформації для прийняття наступних рішень ботом.",
        ),
    ]
    for title_text, code, explanation in python_blocks:
        add_heading(doc, title_text, level=3)
        add_code_block(doc, code)
        add_paragraph(doc, explanation)

    add_heading(doc, "5.8. Висновок по ЛР 2", level=2)
    add_paragraph(
        doc,
        "У проєкті SQL-запити є частиною робочої логіки системи: одні читають аналітичні дані, інші записують прогнози, "
        "треті оновлюють або замінюють моделі.",
    )

    add_heading(doc, "6. Лабораторна робота №3. Підпрограми СУБД", level=1)
    add_heading(doc, "6.1. Що саме з ЛР 3 використано у проєкті", level=2)
    add_paragraph(
        doc,
        "У третій лабораторній роботі вивчаються підпрограми СУБД: процедури та функції, які виконують логіку безпосередньо всередині бази даних. "
        "У проєкті Crypto Pulse цю роль виконують функції PL/pgSQL.",
    )
    add_heading(doc, "6.2. Функція handle_new_user()", level=2)
    add_code_block(doc, "CREATE OR REPLACE FUNCTION public.handle_new_user()\nRETURNS trigger\nLANGUAGE plpgsql")
    add_paragraph(
        doc,
        "CREATE OR REPLACE FUNCTION створює нову функцію або оновлює вже наявну. "
        "Оскільки вона повертає trigger, цю функцію можна викликати тригером. Вказівка LANGUAGE plpgsql означає, "
        "що логіка функції написана мовою процедурного розширення PostgreSQL.",
    )
    add_paragraph(
        doc,
        "Функція автоматично створює профіль користувача в public.users після реєстрації. Вона переносить дані з auth.users, "
        "встановлює значення за замовчуванням і через ON CONFLICT (id) DO UPDATE оновлює профіль, якщо запис уже існує.",
    )
    for item in [
        "береться новий запис із auth.users;",
        "із raw_user_meta_data читаються username, first_name, last_name, phone_number, birth_date та region;",
        "створюється запис у public.users;",
        "якщо профіль уже є, оновлюються потрібні поля.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.3. Функція check_email_exists(lookup_email text)", level=2)
    add_code_block(doc, "CREATE OR REPLACE FUNCTION public.check_email_exists(lookup_email text)\nRETURNS boolean")
    add_paragraph(
        doc,
        "У цьому блоці створюється функція з одним вхідним параметром lookup_email типу text. "
        "Оскільки функція повертає boolean, результатом її роботи може бути лише true або false.",
    )
    add_code_block(
        doc,
        "RETURN EXISTS (\n"
        "  SELECT 1\n"
        "  FROM auth.users\n"
        "  WHERE lower(email) = lower(lookup_email)\n"
        ");",
    )
    add_paragraph(
        doc,
        "RETURN EXISTS перевіряє сам факт наявності хоча б одного відповідного запису. "
        "Функція lower() приводить обидва значення email до нижнього регістру, тому порівняння не залежить від великих і малих літер.",
    )
    add_paragraph(
        doc,
        "Функція повертає true або false залежно від того, чи вже існує користувач із таким email. "
        "Використання lower() робить перевірку нечутливою до регістру.",
    )
    add_heading(doc, "6.4. Функція delete_current_user()", level=2)
    add_code_block(doc, "CREATE OR REPLACE FUNCTION public.delete_current_user()\nRETURNS void")
    add_paragraph(
        doc,
        "Ця функція має тип повернення void, тобто не передає значення назовні, "
        "а виконує дію - видаляє поточного користувача.",
    )
    add_code_block(doc, "current_user_id := auth.uid();")
    add_paragraph(
        doc,
        "Цей рядок записує в змінну current_user_id id поточного авторизованого користувача. "
        "Завдяки цьому функція працює саме з тим користувачем, який викликав її, а не з довільним id.",
    )
    add_paragraph(
        doc,
        "Функція видаляє саме поточного авторизованого користувача. Вона не приймає id ззовні, а бере його через auth.uid(), "
        "тому користувач не може випадково або навмисно видалити чужий акаунт. Після видалення запису з auth.users пов'язані дані "
        "очищуються завдяки ON DELETE CASCADE.",
    )
    add_heading(doc, "6.5. Де це видно в проєкті", level=2)
    add_paragraph(
        doc,
        "Ці функції розміщені в SQL-частині проєкту і відповідають за логіку користувачів. Python-код окремо не дублює цю логіку: "
        "створення профілю, перевірка email і видалення акаунта винесені на рівень СУБД.",
    )
    add_heading(doc, "6.6. Висновок по ЛР 3", level=2)
    add_paragraph(
        doc,
        "У проєкті функції СУБД використовуються для автоматизації повторюваних дій і забезпечення єдиного способу обробки даних користувача.",
    )

    add_heading(doc, "7. Лабораторна робота №4. Тригери", level=1)
    add_heading(doc, "7.1. Що саме з ЛР 4 використано у проєкті", level=2)
    add_paragraph(
        doc,
        "У четвертій лабораторній роботі вивчаються тригери. У проєкті Crypto Pulse використано тригер, який автоматизує створення профілю користувача.",
    )
    add_heading(doc, "7.2. Тригер on_auth_user_created", level=2)
    add_code_block(
        doc,
        "CREATE TRIGGER on_auth_user_created\n"
        "AFTER INSERT ON auth.users\n"
        "FOR EACH ROW\n"
        "EXECUTE FUNCTION public.handle_new_user();",
    )
    add_paragraph(
        doc,
        "Оператор CREATE TRIGGER створює тригер on_auth_user_created. "
        "Він прив'язується до таблиці auth.users, спрацьовує після додавання нового рядка і автоматично викликає функцію handle_new_user(), "
        "яка створює профіль у public.users.",
    )
    for item in [
        "AFTER INSERT означає запуск після створення нового користувача;",
        "ON auth.users означає, що відстежується системна таблиця Supabase Auth;",
        "FOR EACH ROW означає окреме спрацювання для кожного нового користувача;",
        "EXECUTE FUNCTION викликає handle_new_user(), яка створює профіль.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "7.3. Як це працює в проєкті", level=2)
    for item in [
        "Користувач реєструється в застосунку.",
        "Supabase створює новий запис у auth.users.",
        "Тригер on_auth_user_created спрацьовує автоматично.",
        "Викликається функція handle_new_user().",
        "У public.users створюється профіль користувача.",
    ]:
        add_numbered(doc, item)
    add_paragraph(
        doc,
        "Завдяки цьому frontend і backend не повинні окремо виконувати ще один запит для створення профілю: база даних робить це автоматично.",
    )
    add_heading(doc, "7.4. Висновок по ЛР 4", level=2)
    add_paragraph(
        doc,
        "У проєкті тригер використано для реального бізнес-процесу реєстрації, тобто механізм тригерів застосовано на практиці.",
    )

    add_heading(doc, "8. Лабораторна робота №5. Адміністрування БД", level=1)
    add_heading(doc, "8.1. Що саме з ЛР 5 використано у проєкті", level=2)
    add_paragraph(
        doc,
        "У п'ятій лабораторній роботі розглядаються користувачі, ролі та привілеї. У Crypto Pulse це реалізовано через ролі Supabase, "
        "команди GRANT і Row-Level Security.",
    )
    add_heading(doc, "8.2. Приклади надання прав", level=2)
    add_code_block(
        doc,
        "GRANT SELECT, INSERT, UPDATE, DELETE ON public.user_favorites TO authenticated;\n"
        "GRANT SELECT, INSERT, UPDATE, DELETE ON public.alerts TO authenticated;\n"
        "GRANT SELECT, INSERT, UPDATE ON public.users TO authenticated;\n"
        "GRANT EXECUTE ON FUNCTION public.check_email_exists(text) TO anon, authenticated;\n"
        "GRANT EXECUTE ON FUNCTION public.delete_current_user() TO authenticated;",
    )
    add_paragraph(
        doc,
        "Оператор GRANT надає ролям конкретні права. "
        "У цьому прикладі authenticated може працювати з потрібними таблицями, anon отримує доступ лише до перевірки email, "
        "а видалення акаунта залишається доступним тільки авторизованому користувачу.",
    )
    for item in [
        "authenticated може працювати зі своїми обраними монетами та сповіщеннями;",
        "authenticated може читати й редагувати свій профіль;",
        "anon може перевірити email до завершення реєстрації;",
        "видалення акаунта дозволено лише авторизованому користувачу.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.3. Row-Level Security для user_favorites", level=2)
    add_code_block(
        doc,
        "ALTER TABLE public.user_favorites ENABLE ROW LEVEL SECURITY;\n\n"
        "CREATE POLICY \"Users can read own favorites\"\n"
        "ON public.user_favorites\n"
        "FOR SELECT\n"
        "TO authenticated\n"
        "USING (auth.uid() = user_id);",
    )
    add_paragraph(
        doc,
        "ALTER TABLE вмикає Row-Level Security для user_favorites, а CREATE POLICY задає конкретне правило читання. "
        "Умова auth.uid() = user_id дозволяє користувачу бачити лише ті рядки, які належать саме йому.",
    )
    add_paragraph(
        doc,
        "GRANT дає право працювати з таблицею, а RLS визначає, з якими саме рядками це дозволено робити. "
        "Умова auth.uid() = user_id означає, що користувач бачить лише власні записи.",
    )
    add_heading(doc, "8.4. Row-Level Security для alerts", level=2)
    add_code_block(
        doc,
        "CREATE POLICY \"Users can read own alerts\"\n"
        "ON public.alerts\n"
        "FOR SELECT\n"
        "TO authenticated\n"
        "USING (auth.uid() = user_id);",
    )
    add_paragraph(
        doc,
        "Тут створюється окрема політика доступу для таблиці alerts. "
        "Вона працює за тим самим принципом, що й політика для favorites: кожен користувач має доступ лише до власних сповіщень.",
    )
    add_paragraph(
        doc,
        "Така політика потрібна, тому що сповіщення є персональними налаштуваннями користувача і не повинні бути доступні іншим.",
    )
    add_heading(doc, "8.5. Row-Level Security для users", level=2)
    add_code_block(
        doc,
        "CREATE POLICY \"Users can read own profile\"\n"
        "ON public.users\n"
        "FOR SELECT\n"
        "TO authenticated\n"
        "USING (auth.uid() = id);",
    )
    add_paragraph(
        doc,
        "У таблиці users порівняння виконується з полем id, бо саме воно є ключем профілю. "
        "Отже, користувач може прочитати тільки власний профіль.",
    )
    add_paragraph(
        doc,
        "У таблиці users порівняння виконується з полем id, оскільки саме воно є ідентифікатором профілю.",
    )
    add_heading(doc, "8.6. Де це видно в логіці проєкту", level=2)
    add_paragraph(
        doc,
        "SQL-рівень доступу напряму пов'язаний із логікою застосунку: user_favorites містить персональний список монет, "
        "alerts містить приватні сповіщення, users містить особисті дані профілю. Саме тому для цих таблиць застосовано окремі політики доступу.",
    )
    add_heading(doc, "8.7. Заплановані рівні підписки та обмеження доступу", level=2)
    add_paragraph(
        doc,
        "У проєкті вже передбачені поля, пов'язані з підпискою користувача.",
    )
    add_code_block(
        doc,
        "subscription VARCHAR(20) DEFAULT 'free',\n"
        "active_plan VARCHAR(20) DEFAULT 'free',\n"
        "billing_cycle VARCHAR(10) DEFAULT 'monthly'",
    )
    add_paragraph(
        doc,
        "Ці поля зберігаються в таблиці public.users і вже зараз дають основу для поділу користувачів за рівнями доступу. "
        "Поле subscription може описувати загальний тип підписки, active_plan - фактично активний тариф, а billing_cycle - спосіб оплати, "
        "наприклад щомісячний або річний.",
    )
    add_paragraph(doc, "Надалі на сайті планується реалізувати кілька рівнів підписки:")
    for item in [
        "free - базовий тариф;",
        "pro - розширений тариф;",
        "premium - повний доступ до всіх можливостей.",
    ]:
        add_numbered(doc, item)
    add_paragraph(doc, "Можлива логіка обмежень:")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 7160])
    for idx, text in enumerate(["Рівень підписки", "Можливості"]):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        style_run(cell.paragraphs[0].add_run(text), bold=True)
    subscription_rows = [
        ("free", "Обмежена кількість обраних монет, невелика кількість alerts, базові прогнози."),
        ("pro", "Більше обраних монет, більше alerts, доступ до розширеної статистики та частіших оновлень."),
        ("premium", "Максимальні ліміти, повний доступ до аналітики, усі таймфрейми та пріоритетні функції."),
    ]
    for row in subscription_rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = ""
            style_run(cells[idx].paragraphs[0].add_run(text))
    add_paragraph(
        doc,
        "Таку систему планується реалізувати у два рівні:",
    )
    for item in [
        "На рівні бази даних - зберігати активний план користувача в public.users, а за потреби додати окрему таблицю subscription_plans, де будуть описані ліміти кожного тарифу.",
        "На рівні логіки застосунку - перед виконанням дії перевіряти active_plan користувача та порівнювати його з дозволеними лімітами.",
    ]:
        add_numbered(doc, item)
    add_paragraph(
        doc,
        "Наприклад, для тарифів можна передбачити максимальну кількість favorites, максимальну кількість alerts, доступні таймфрейми та доступ до розширеної аналітики.",
    )
    add_code_block(
        doc,
        "CREATE TABLE public.subscription_plans (\n"
        "  code VARCHAR(20) PRIMARY KEY,\n"
        "  max_favorites INTEGER NOT NULL,\n"
        "  max_alerts INTEGER NOT NULL,\n"
        "  allow_advanced_analytics BOOLEAN NOT NULL DEFAULT false,\n"
        "  allow_all_timeframes BOOLEAN NOT NULL DEFAULT false\n"
        ");",
    )
    add_paragraph(
        doc,
        "Оператор CREATE TABLE у цьому прикладі створює довідник тарифних планів. "
        "Поле code містить назву тарифу, а інші стовпці задають конкретні ліміти та можливості кожного плану. "
        "Такий підхід дозволяє змінювати правила тарифів без зміни основної структури таблиці користувачів.",
    )
    add_paragraph(
        doc,
        "Наприклад, для перевірки кількості alerts можна буде використати такий запит:",
    )
    add_code_block(
        doc,
        "SELECT COUNT(*)\n"
        "FROM public.alerts\n"
        "WHERE user_id = auth.uid();",
    )
    add_paragraph(
        doc,
        "Цей SELECT рахує, скільки сповіщень уже створив поточний користувач. "
        "Після цього кількість можна порівняти з лімітом його тарифу і дозволити або заборонити створення нового alert.",
    )
    add_paragraph(
        doc,
        "Таким чином рівні підписки надалі будуть використовуватися не лише як інформаційні поля профілю, а як основа для реального розмежування функціональності сайту.",
    )
    add_heading(doc, "8.8. Висновок по ЛР 5", level=2)
    add_paragraph(
        doc,
        "У проєкті адміністрування БД реалізоване через ролі, привілеї та політики безпеки. Особливо важливим є Row-Level Security, "
        "який контролює доступ не лише до таблиці загалом, а й до конкретних рядків.",
    )

    add_heading(doc, "9. Загальний висновок", level=1)
    add_paragraph(
        doc,
        "У проєкті Crypto Pulse реально використовуються основні теми лабораторних робіт з баз даних: побудова схеми реляційної БД, "
        "ключі, зв'язки, обмеження цілісності, SQL-запити різних типів, функції СУБД, тригери, ролі, привілеї та політики доступу.",
    )
    add_paragraph(
        doc,
        "База даних у проєкті не є окремим формальним елементом, а безпосередньо бере участь у роботі застосунку: "
        "зберігає користувачів, сповіщення, обрані монети, ML-моделі та результати прогнозів.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
